# Copyright (C) 2026 Patrick R. Wallace <mail.prw@gmail.com>
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program. If not, see <https://www.gnu.org/licenses/>.

"""doc-scanner.py

Improved doc-scanner with argparse, regex support, exclusion patterns,
logging/verbosity, and optional parallel scanning.

Features
- Search `.txt`, `.docx`, and optional `.doc` files (if extractor available).
- Case-insensitive substring matching by default, optional regex via `--regex`.
- Exclude files or paths using glob patterns via `--exclude` (repeatable).
- Verbosity with `-v`/`--verbose` and logging.
- Parallel scanning with `--workers` (uses ThreadPoolExecutor).

Usage:
    python doc-scanner.py <directory> <search_string> [options]

See --help for full options.
"""

import argparse
import fnmatch
import logging
import os
import re
import shutil
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
except Exception:
    Document = None
    PackageNotFoundError = Exception

# Optional .doc support: uncommon module `python_doc` or fallback to `textract` if available
try:
    import python_doc as doc  # optional/legacy
except Exception:
    doc = None

try:
    import textract  # optional fallback for .doc extraction
except Exception:
    textract = None

def search_txt(file_path, matcher):
    """Search a .txt file using matcher(text) -> bool. Return list of 1-based line numbers."""
    results = []
    try:
        with open(file_path, "r", errors="ignore") as f:
            for i, line in enumerate(f, 1):
                try:
                    if matcher(line):
                        results.append(i)
                except Exception:
                    # Ignore malformed lines that crash matcher
                    continue
    except Exception as e:
        logging.warning("Failed to read text file %s: %s", file_path, e)
    return results

def search_docx(file_path, matcher):
    """Search a .docx file using matcher(text) -> bool. Return list of 1-based paragraph indices."""
    results = []
    if Document is None:
        logging.warning("python-docx not available; skipping .docx file %s", file_path)
        return results
    try:
        document = Document(file_path)
        for i, para in enumerate(document.paragraphs, 1):
            try:
                if matcher(para.text):
                    results.append(i)
            except Exception:
                continue
    except PackageNotFoundError:
        logging.debug("Not a valid docx package: %s", file_path)
    except Exception as e:
        logging.warning("Failed to read .docx file %s: %s", file_path, e)
    return results

def search_doc(file_path, matcher):
    """Search a .doc file using matcher(text) -> bool. Return list of 1-based paragraph indices.

    Uses `python_doc` if installed, otherwise attempts `textract` (if available).
    If neither is present the function returns an empty list and a warning is logged.
    """
    results = []
    if doc is not None:
        try:
            d = doc.Document(file_path)
            for i, para in enumerate(d.paragraphs, 1):
                try:
                    if matcher(para.text):
                        results.append(i)
                except Exception:
                    continue
        except Exception as e:
            logging.debug("python_doc failed to read %s: %s", file_path, e)
        return results

    if textract is not None:
        try:
            raw = textract.process(file_path)
            try:
                text = raw.decode("utf-8", errors="ignore")
            except Exception:
                text = str(raw)
            for i, para in enumerate(text.splitlines(), 1):
                try:
                    if matcher(para):
                        results.append(i)
                except Exception:
                    continue
        except Exception as e:
            logging.debug("textract failed to read %s: %s", file_path, e)
        return results

    logging.info("No .doc extractor installed; skipping .doc file: %s", file_path)
    return results

def _should_exclude(rel_path, name, exclude_patterns):
    if not exclude_patterns:
        return False
    parts = rel_path.split(os.sep)
    for pat in exclude_patterns:
        if fnmatch.fnmatch(rel_path, pat) or fnmatch.fnmatch(name, pat):
            return True
        if any(fnmatch.fnmatch(p, pat) or p == pat for p in parts):
            return True
    return False


def scan_directory(directory, matcher, exclude_patterns=None, workers=1):
    """Scan `directory` for supported files and apply `matcher` to each file.

    Returns (results, matched_files).
    - `matcher` is a callable that accepts a single text argument and returns True/False.
    - `exclude_patterns` is a list of glob patterns to skip.
    - `workers` > 1 enables threaded parallel scanning.
    """
    candidate_files = []
    directory = os.path.abspath(directory)
    for root, dirs, files in os.walk(directory):
        # Prune directories that match exclude patterns
        for d in list(dirs):
            rel_dir = os.path.relpath(os.path.join(root, d), directory)
            if _should_exclude(rel_dir, d, exclude_patterns):
                logging.debug("Skipping directory %s due to exclude pattern", rel_dir)
                dirs.remove(d)
        for file in files:
            name = file
            rel_path = os.path.relpath(os.path.join(root, file), directory)
            if _should_exclude(rel_path, name, exclude_patterns):
                logging.debug("Excluding %s", rel_path)
                continue
            if not name.lower().endswith((".txt", ".docx", ".doc")):
                continue
            candidate_files.append(os.path.join(root, file))

    results = []
    matched_files = []

    def _process(file_path):
        name = file_path.lower()
        try:
            if name.endswith(".txt"):
                lines = search_txt(file_path, matcher)
            elif name.endswith(".docx"):
                lines = search_docx(file_path, matcher)
            elif name.endswith(".doc"):
                lines = search_doc(file_path, matcher)
            else:
                lines = []
            return file_path, lines
        except Exception as e:
            logging.debug("Error processing %s: %s", file_path, e)
            return file_path, []

    if workers and workers > 1:
        logging.info("Scanning with %d worker threads", workers)
        with ThreadPoolExecutor(max_workers=workers) as exe:
            future_to_file = {exe.submit(_process, f): f for f in candidate_files}
            for fut in as_completed(future_to_file):
                file_path, lines = fut.result()
                if lines:
                    results.append(f"{file_path}: {lines}")
                    matched_files.append(file_path)
    else:
        for f in candidate_files:
            file_path, lines = _process(f)
            if lines:
                results.append(f"{file_path}: {lines}")
                matched_files.append(file_path)

    return results, matched_files

def _setup_logging(verbosity: int):
    if verbosity >= 2:
        level = logging.DEBUG
    elif verbosity == 1:
        level = logging.INFO
    else:
        level = logging.WARNING
    logging.basicConfig(format="%(levelname)s: %(message)s", level=level)


def main(argv=None):
    parser = argparse.ArgumentParser(description="Search .txt/.docx/.doc files for a string or regex pattern.")
    parser.add_argument("directory", help="Root folder to scan (recursive)")
    parser.add_argument("search_string", help="Substring (default) or regex pattern to search for")
    parser.add_argument("-o", "--output", help="Write plain-text results to the specified file (overwrites)")
    parser.add_argument("-c", "--copy", help="Copy matched files to an existing folder (uses shutil.copy2)")
    parser.add_argument("-r", "--regex", action="store_true", help="Treat search_string as a regular expression")
    parser.add_argument("-e", "--exclude", action="append", default=[], help="Glob pattern to exclude (repeatable)")
    parser.add_argument("-v", "--verbose", action="count", default=0, help="Increase verbosity (use -v multiple times)")
    parser.add_argument("-w", "--workers", type=int, default=1, help="Number of parallel worker threads (1 disables parallelism)")

    args = parser.parse_args(argv)

    _setup_logging(args.verbose)

    directory = args.directory
    search_string = args.search_string
    output_file = args.output
    copy_folder = args.copy
    use_regex = args.regex
    exclude_patterns = args.exclude
    workers = args.workers

    if not os.path.isdir(directory):
        logging.error("Directory does not exist: %s", directory)
        sys.exit(1)

    if copy_folder and not os.path.isdir(copy_folder):
        logging.error("Copy destination does not exist: %s", copy_folder)
        sys.exit(1)

    # Build matcher
    if use_regex:
        try:
            pattern = re.compile(search_string, flags=re.IGNORECASE)
            matcher = lambda text: bool(pattern.search(text or ""))
        except re.error as e:
            logging.error("Invalid regular expression: %s", e)
            sys.exit(1)
    else:
        needle = search_string.lower()
        matcher = lambda text: needle in (text or "").lower()

    results, matched_files = scan_directory(directory, matcher, exclude_patterns=exclude_patterns, workers=workers)

    if output_file:
        try:
            with open(output_file, "w", encoding="utf-8") as f:
                for line in results:
                    f.write(line + "\n")
            logging.info("Results written to %s", output_file)
        except Exception as e:
            logging.error("Failed to write output file %s: %s", output_file, e)
    else:
        for line in results:
            print(line)

    if copy_folder:
        for file_path in matched_files:
            try:
                shutil.copy2(file_path, copy_folder)
                logging.info("Copied: %s -> %s", file_path, copy_folder)
            except Exception as e:
                logging.error("Failed to copy %s: %s", file_path, e)


if __name__ == "__main__":
    main()

